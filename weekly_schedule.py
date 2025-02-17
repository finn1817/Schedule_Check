import pandas as pd
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from docx import Document
from PIL import Image, ImageDraw, ImageFont
import random
import time


class WeeklyScheduleGenerator:
    def __init__(self, root):
        """Initialize main GUI components."""
        self.root = root
        self.root.title("Weekly Schedule Generator")
        self.root.geometry("500x400")
        
        # initializing UI elements
        main_frame = ttk.Frame(root, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)

        title_label = ttk.Label(main_frame, text="Weekly Schedule Generator", font=('Helvetica', 16, 'bold'))
        title_label.pack(pady=(0, 20))

        self.load_schedule_btn = ttk.Button(main_frame, text="Load Excel Schedule", command=self.load_schedule)
        self.load_schedule_btn.pack(pady=10, fill=tk.X)

        self.generate_btn = ttk.Button(main_frame, text="Generate Schedule (Hourly)", command=self.generate_schedule)
        self.generate_btn.pack(pady=10, fill=tk.X)

        self.save_btn = ttk.Button(main_frame, text="Save Word File", command=self.save_word_file)
        self.save_btn.pack(pady=10, fill=tk.X)

        self.image_path = "hourly_schedule.png"
        self.schedule_df = None
        self.schedule_data = None

    def load_schedule(self):
        """Load the Excel schedule file."""
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
        if not file_path:
            return

        try:
            self.schedule_df = pd.read_excel(file_path)
            required_columns = {'First Name', 'Last Name', 'Email', 'Sunday', 'Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday'}
            if not required_columns.issubset(self.schedule_df.columns):
                messagebox.showerror("Error", f"Excel must contain these columns: {', '.join(required_columns)}")
            else:
                messagebox.showinfo("Success", "Schedule file loaded successfully!")

        except Exception as e:
            messagebox.showerror("Error", f"Error loading file: {e}")

    def generate_schedule(self):
        """Generate an hourly schedule based on loaded data."""
        if self.schedule_df is None:
            messagebox.showerror("Error", "Please load a schedule file first.")
            return

        # start an empty schedule
        hour_schedule = {day: {hour: None for hour in range(8, 24)} for day in ['Sunday', 'Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday']}

        # assign workers to slots
        for _, row in self.schedule_df.iterrows():
            for day in hour_schedule.keys():
                if pd.notna(row[day]) and row[day].lower() != 'na':
                    available_hours = range(8, 24)  # ex) availability from 8 AM to 11 PM
                    for hour in available_hours:
                        if hour_schedule[day][hour] is None:
                            hour_schedule[day][hour] = f"{row['First Name']} {row['Last Name']}"
                            break

        self.schedule_data = hour_schedule
        self.create_visual_schedule()

        messagebox.showinfo("Success", "Hourly schedule has been generated!")

    def create_visual_schedule(self):
        """Generate a visual PNG image for the hourly schedule."""
        width, height = 1200, 800
        img = Image.new("RGB", (width, height), (255, 255, 255))
        draw = ImageDraw.Draw(img)

        try:
            font_title = ImageFont.truetype("arial.ttf", 24)
            font = ImageFont.truetype("arial.ttf", 20)
        except:
            font_title = font = ImageFont.load_default()

        # header
        draw.text((10, 10), "Hourly Schedule", fill=(0, 0, 0), font=font_title)

        y = 50
        for day, hours in self.schedule_data.items():
            draw.text((10, y), day, fill=(0, 0, 0), font=font)
            for hour, worker in hours.items():
                shift_text = f"{hour}:00 - {hour + 1}:00: {worker if worker else 'Unassigned'}"
                y += 25
                draw.text((50, y), shift_text, fill=(0, 0, 0), font=font)
            y += 50

        img.save(self.image_path)

    def save_word_file(self):
        """Save the schedule as a Word document."""
        if self.schedule_data is None:
            messagebox.showerror("Error", "Please generate a schedule first.")
            return

        try:
            doc = Document()
            doc.add_heading("Hourly Schedule", level=1)

            for day, hours in self.schedule_data.items():
                doc.add_heading(day, level=2)
                for hour, worker in hours.items():
                    doc.add_paragraph(f"{hour}:00 - {hour + 1}:00: {worker if worker else 'Unassigned'}")

            file_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word files", "*.docx")]
            )
            if file_path:
                doc.save(file_path)
                messagebox.showinfo("Success", "Schedule saved successfully!")
        except Exception as e:
            messagebox.showerror("Error", f"Error saving Word file: {e}")


if __name__ == "__main__":
    root = tk.Tk()
    app = WeeklyScheduleGenerator(root)
    root.mainloop()
